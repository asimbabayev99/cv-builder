from __future__ import annotations

import io
import os

from jinja2 import Environment, FileSystemLoader
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.cv import Resume

_TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")


class ExportService:
    @staticmethod
    def generate_pdf(resume: Resume) -> bytes:
        env = Environment(loader=FileSystemLoader(_TEMPLATE_DIR))
        template = env.get_template("resume.html")
        html = template.render(resume=resume, color_hex=resume.color_hex or "#111118")

        buf = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.StringIO(html), dest=buf)
        if pisa_status.err:
            raise RuntimeError("PDF generation failed")
        return buf.getvalue()

    @staticmethod
    def generate_docx(resume: Resume) -> bytes:
        doc = Document()

        # -- Helper to parse hex color --------------------------------
        def _color(hex_str: str | None) -> RGBColor:
            h = (hex_str or "#111118").lstrip("#")
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        accent = _color(resume.color_hex)

        # -- Name heading ---------------------------------------------
        name = f"{resume.first_name or ''} {resume.last_name or ''}".strip()
        if name:
            h = doc.add_heading(name, level=0)
            for run in h.runs:
                run.font.color.rgb = accent

        # -- Professional title ----------------------------------------
        if resume.professional_title:
            p = doc.add_paragraph()
            run = p.add_run(resume.professional_title)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # -- Contact info ----------------------------------------------
        contact_parts = [
            v for v in [resume.email, resume.phone, resume.location] if v
        ]
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        # -- Summary ---------------------------------------------------
        if resume.summary:
            h = doc.add_heading("Summary", level=1)
            for run in h.runs:
                run.font.color.rgb = accent
            doc.add_paragraph(resume.summary)

        # -- Experience ------------------------------------------------
        if resume.experiences:
            h = doc.add_heading("Experience", level=1)
            for run in h.runs:
                run.font.color.rgb = accent
            for exp in resume.experiences:
                p = doc.add_paragraph()
                title_run = p.add_run(exp.job_title)
                title_run.bold = True
                if exp.company:
                    p.add_run(f" — {exp.company}")
                date_parts = []
                if exp.start_date:
                    date_parts.append(exp.start_date)
                if exp.currently_working:
                    date_parts.append("Present")
                elif exp.end_date:
                    date_parts.append(exp.end_date)
                if date_parts:
                    dp = doc.add_paragraph(" – ".join(date_parts))
                    dp.paragraph_format.space_before = Pt(0)
                    for run in dp.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                if exp.description:
                    doc.add_paragraph(exp.description)

        # -- Education -------------------------------------------------
        if resume.educations:
            h = doc.add_heading("Education", level=1)
            for run in h.runs:
                run.font.color.rgb = accent
            for edu in resume.educations:
                p = doc.add_paragraph()
                inst_run = p.add_run(edu.institution)
                inst_run.bold = True
                if edu.degree or edu.field_of_study:
                    parts = [v for v in [edu.degree, edu.field_of_study] if v]
                    doc.add_paragraph(" in ".join(parts))
                date_parts = []
                if edu.start_date:
                    date_parts.append(edu.start_date)
                if edu.currently_studying:
                    date_parts.append("Present")
                elif edu.end_date:
                    date_parts.append(edu.end_date)
                if date_parts:
                    dp = doc.add_paragraph(" – ".join(date_parts))
                    for run in dp.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                if edu.description:
                    doc.add_paragraph(edu.description)

        # -- Skills ----------------------------------------------------
        if resume.skills:
            h = doc.add_heading("Skills", level=1)
            for run in h.runs:
                run.font.color.rgb = accent
            technical = [s for s in resume.skills if s.category == "technical"]
            soft = [s for s in resume.skills if s.category == "soft"]
            if technical:
                p = doc.add_paragraph()
                p.add_run("Technical: ").bold = True
                p.add_run(", ".join(s.name for s in technical))
            if soft:
                p = doc.add_paragraph()
                p.add_run("Soft Skills: ").bold = True
                p.add_run(", ".join(s.name for s in soft))

        # -- Languages -------------------------------------------------
        if resume.languages:
            h = doc.add_heading("Languages", level=1)
            for run in h.runs:
                run.font.color.rgb = accent
            for lang in resume.languages:
                doc.add_paragraph(f"{lang.name} — {lang.proficiency}")

        # -- Certificates ----------------------------------------------
        if resume.certificates:
            h = doc.add_heading("Certificates", level=1)
            for run in h.runs:
                run.font.color.rgb = accent
            for cert in resume.certificates:
                p = doc.add_paragraph()
                p.add_run(cert.name).bold = True
                if cert.organization:
                    doc.add_paragraph(cert.organization)
                date_info = []
                if cert.issue_date:
                    date_info.append(f"Issued: {cert.issue_date}")
                if cert.expiration_date:
                    date_info.append(f"Expires: {cert.expiration_date}")
                if cert.no_expiry:
                    date_info.append("No expiration")
                if date_info:
                    dp = doc.add_paragraph(" | ".join(date_info))
                    for run in dp.runs:
                        run.font.size = Pt(9)
                if cert.description:
                    doc.add_paragraph(cert.description)

        # -- Custom Sections -------------------------------------------
        if resume.custom_sections:
            for section in resume.custom_sections:
                h = doc.add_heading(section.title, level=1)
                for run in h.runs:
                    run.font.color.rgb = accent
                if section.description:
                    doc.add_paragraph(section.description)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
